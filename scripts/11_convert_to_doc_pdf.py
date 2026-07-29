"""
Script 11 — Convert paper_draft.md into publication-grade .docx (Word) and .pdf formats.

Usage:
    python scripts/11_convert_to_doc_pdf.py
"""
import re
from pathlib import Path

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable


def parse_markdown(md_text):
    """Parses markdown lines into structured blocks (heading, paragraph, table, hr)."""
    lines = md_text.splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        
        # Heading
        if line.startswith("#"):
            level = len(line.split()[0])
            text = line.lstrip("#").strip()
            blocks.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        # Horizontal Rule
        if line in ["---", "***", "___"]:
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Markdown Table
        if "|" in line and i + 1 < len(lines) and "|---" in lines[i + 1]:
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            
            headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
            rows = []
            for t_line in table_lines[2:]:
                cols = [c.strip() for c in t_line.split("|")[1:-1]]
                if cols:
                    rows.append(cols)
            blocks.append({"type": "table", "headers": headers, "rows": rows})
            continue

        # Codeblock or pre
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1 # skip ending ```
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            continue

        # Bullet List
        if line.startswith("- ") or line.startswith("* "):
            blocks.append({"type": "bullet", "text": line[2:].strip()})
            i += 1
            continue

        # Standard Paragraph
        blocks.append({"type": "paragraph", "text": line})
        i += 1
        
    return blocks


def clean_md_inline(text):
    """Removes inline markdown tags like **, *, ` for clean text rendering."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    return text


def build_docx(blocks, output_path):
    """Builds a formatted Word document (.docx)."""
    doc = Document()
    
    # Page Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for b in blocks:
        b_type = b["type"]
        
        if b_type == "heading":
            level = b["level"]
            text = clean_md_inline(b["text"])
            h = doc.add_heading(text, level=min(level, 3))
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if level == 1:
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(6)
            elif level == 2:
                h.paragraph_format.space_before = Pt(10)
                h.paragraph_format.space_after = Pt(4)

        elif b_type == "paragraph":
            text = clean_md_inline(b["text"])
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15

        elif b_type == "bullet":
            text = clean_md_inline(b["text"])
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)

        elif b_type == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

        elif b_type == "table":
            headers = [clean_md_inline(h) for h in b["headers"]]
            rows = [[clean_md_inline(c) for c in r] for r in b["rows"]]
            
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header Row
            hdr_cells = table.rows[0].cells
            for idx, name in enumerate(headers):
                hdr_cells[idx].text = name
                hdr_cells[idx].paragraphs[0].runs[0].font.bold = True

            # Data Rows
            for r_idx, row in enumerate(rows):
                row_cells = table.rows[r_idx + 1].cells
                for c_idx, val in enumerate(row):
                    if c_idx < len(row_cells):
                        row_cells[c_idx].text = val
            
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

        elif b_type == "code":
            p = doc.add_paragraph(b["text"])
            p.paragraph_format.space_after = Pt(6)

    doc.save(output_path)
    print(f"Word document successfully created: {output_path}")


def md_to_pdf_paragraph(text, style):
    """Converts inline markdown (bold/italic) to HTML tags for ReportLab Paragraph."""
    # Convert **bold** to <b>bold</b>
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    # Convert *italic* to <i>italic</i>
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    # Convert `code` to <font name="Courier">\1</font>
    text = re.sub(r'`(.*?)`', r'<font name="Courier">\1</font>', text)
    # Escapes unescaped &
    text = text.replace("& ", "&amp; ")
    return Paragraph(text, style)


def build_pdf(blocks, output_path):
    """Builds a publication-grade PDF document (.pdf) via ReportLab."""
    pdf = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=54, leftMargin=54,
        topMargin=54, bottomMargin=54
    )

    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'PdfTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#1A2B4C"),
        alignment=0,
        spaceAfter=8
    )

    h1_style = ParagraphStyle(
        'PdfH1',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#2C3E50"),
        spaceBefore=12,
        spaceAfter=6
    )

    h2_style = ParagraphStyle(
        'PdfH2',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#34495E"),
        spaceBefore=10,
        spaceAfter=4
    )

    body_style = ParagraphStyle(
        'PdfBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#222222"),
        spaceAfter=4
    )

    bullet_style = ParagraphStyle(
        'PdfBullet',
        parent=body_style,
        leftIndent=12,
        bulletIndent=4,
        spaceAfter=2
    )

    table_header_style = ParagraphStyle(
        'PdfTblHdr',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=11,
        textColor=colors.whitesmoke,
        alignment=1
    )

    table_body_style = ParagraphStyle(
        'PdfTblBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        leading=10,
        textColor=colors.black,
        alignment=1
    )

    story = []

    for b in blocks:
        b_type = b["type"]

        if b_type == "heading":
            level = b["level"]
            text = b["text"]
            if level == 1:
                story.append(md_to_pdf_paragraph(text, title_style))
            elif level == 2:
                story.append(md_to_pdf_paragraph(text, h1_style))
            else:
                story.append(md_to_pdf_paragraph(text, h2_style))

        elif b_type == "paragraph":
            story.append(md_to_pdf_paragraph(b["text"], body_style))

        elif b_type == "bullet":
            story.append(md_to_pdf_paragraph(f"• {b['text']}", bullet_style))

        elif b_type == "hr":
            story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#BDC3C7"), spaceAfter=6))

        elif b_type == "table":
            headers = [md_to_pdf_paragraph(h, table_header_style) for h in b["headers"]]
            table_data = [headers]
            for row in b["rows"]:
                row_cells = [md_to_pdf_paragraph(c, table_body_style) for c in row]
                table_data.append(row_cells)

            t = Table(table_data, hAlign='CENTER')
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2C3E50')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 4),
                ('TOPPADDING', (0, 0), (-1, 0), 4),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#BDC3C7')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8F9FA')]),
            ]))
            story.append(Spacer(1, 4))
            story.append(t)
            story.append(Spacer(1, 6))

        elif b_type == "code":
            code_style = ParagraphStyle(
                'PdfCode',
                parent=body_style,
                fontName='Courier',
                fontSize=8,
                leading=10,
                textColor=colors.HexColor("#2C3E50"),
                backColor=colors.HexColor("#F4F6F7"),
                borderColor=colors.HexColor("#BDC3C7"),
                borderWidth=0.5,
                borderPadding=4,
                spaceAfter=6
            )
            story.append(Paragraph(b["text"].replace("\n", "<br/>"), code_style))

    pdf.build(story)
    print(f"PDF document successfully created: {output_path}")


def main():
    md_file = Path("paper_draft.md")
    if not md_file.exists():
        print("Error: paper_draft.md not found!")
        return

    print("Reading paper_draft.md...")
    md_content = md_file.read_text(encoding="utf-8")
    blocks = parse_markdown(md_content)

    docx_path = Path("paper_draft.docx")
    pdf_path = Path("paper_draft.pdf")

    print("Building Word (.docx) document...")
    build_docx(blocks, docx_path)

    print("Building PDF (.pdf) document...")
    build_pdf(blocks, pdf_path)

    print("Done! Formats generated successfully:")
    print(f"  - Word: {docx_path.resolve()}")
    print(f"  - PDF:  {pdf_path.resolve()}")


if __name__ == "__main__":
    main()
