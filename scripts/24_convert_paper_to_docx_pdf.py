"""
Script 24 — Convert paper_draft.md to Microsoft Word (.docx) and PDF (.pdf) with official author details.

Usage:
    python scripts/24_convert_paper_to_docx_pdf.py
"""
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor

MD_FILE = Path("paper_draft.md")
DOCX_FILE = Path("paper_draft.docx")
PDF_FILE = Path("paper_draft.pdf")


def convert_to_docx():
    doc = Document()
    doc.add_heading("Dual-Stream Semantic-Frequency Fusion Network for Explainable Deepfake Detection and Cross-Domain Generalization", 0)
    p_author = doc.add_paragraph()
    p_author.add_run("Shashank Singh and Deepinder Kaur\n").bold = True
    p_author.add_run("Department of Computer Science & Engineering (Artificial Intelligence)\nKIET Deemed to be University, Ghaziabad, India — July 2026")
    
    with open(MD_FILE, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line_s = line.strip()
        if not line_s or line_s.startswith("# Dual-Stream") or line_s.startswith("**Shashank") or line_s.startswith("Department of"):
            continue

        if line_s.startswith("## "):
            doc.add_heading(line_s.replace("## ", ""), level=1)
        elif line_s.startswith("### "):
            doc.add_heading(line_s.replace("### ", ""), level=2)
        elif line_s.startswith("#### "):
            doc.add_heading(line_s.replace("#### ", ""), level=3)
        elif line_s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line_s[2:])
        elif not line_s.startswith("|") and not line_s.startswith("![") and not line_s.startswith("```"):
            doc.add_paragraph(line_s)

    doc.save(DOCX_FILE)
    print(f"Saved: {DOCX_FILE}")


def convert_to_pdf():
    doc = SimpleDocTemplate(str(PDF_FILE), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle("TitleStyle", parent=styles["Title"], fontSize=15, leading=19, textColor=HexColor("#1A2B4C"))
    h1_style = ParagraphStyle("H1Style", parent=styles["Heading1"], fontSize=12, leading=15, textColor=HexColor("#27AE60"))
    body_style = ParagraphStyle("BodyStyle", parent=styles["BodyText"], fontSize=9.5, leading=13.5)

    story.append(Paragraph("Dual-Stream Semantic-Frequency Fusion Network for Explainable Deepfake Detection and Cross-Domain Generalization", title_style))
    story.append(Spacer(1, 8))
    story.append(Paragraph("<b>Shashank Singh</b> and <b>Deepinder Kaur</b><br/>Department of Computer Science & Engineering (Artificial Intelligence)<br/>KIET Deemed to be University, Ghaziabad, India — July 2026", body_style))
    story.append(Spacer(1, 12))

    with open(MD_FILE, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line_s = line.strip()
        if not line_s or line_s.startswith("# Dual-Stream") or line_s.startswith("**Shashank") or line_s.startswith("Department of"):
            continue

        if line_s.startswith("## "):
            story.append(Spacer(1, 8))
            story.append(Paragraph(line_s.replace("## ", ""), h1_style))
            story.append(Spacer(1, 4))
        elif not line_s.startswith("|") and not line_s.startswith("![") and not line_s.startswith("```"):
            clean_text = line_s.replace("*", "").replace("#", "")
            if clean_text:
                story.append(Paragraph(clean_text, body_style))
                story.append(Spacer(1, 3))

    doc.build(story)
    print(f"Saved: {PDF_FILE}")


if __name__ == "__main__":
    convert_to_docx()
    convert_to_pdf()
